import os
import PyPDF2
import docx
from werkzeug.datastructures import FileStorage

def extract_text_from_resume(file):
    """Extract text from PDF and DOCX (supports file path OR Flask FileStorage)"""

    text = ""

    # ✅ CASE 1: FileStorage (uploaded file)
    if isinstance(file, FileStorage):
        filename = file.filename.lower()

        if filename.endswith(".pdf"):
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() or ""

        elif filename.endswith(".docx"):
            doc = docx.Document(file)
            text = "\n".join([p.text for p in doc.paragraphs])

        elif filename.endswith(".txt"):
            text = file.read().decode("utf-8")

        else:
            raise ValueError("Unsupported file type")

    # ✅ CASE 2: File path (string)
    elif isinstance(file, (str, os.PathLike)):
        file_extension = os.path.splitext(file)[1].lower()

        if file_extension == ".pdf":
            with open(file, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() or ""

        elif file_extension == ".docx":
            doc = docx.Document(file)
            text = "\n".join([p.text for p in doc.paragraphs])

        elif file_extension == ".txt":
            with open(file, "r", encoding="utf-8") as f:
                text = f.read()

        else:
            raise ValueError("Unsupported file type")

    else:
        raise TypeError("Invalid file type")

    return text.strip()